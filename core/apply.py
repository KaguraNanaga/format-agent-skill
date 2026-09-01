# 执行器接线（PLAN.md 第 7 节）：
# apply_format(docx_path, spec, rolemap, out_path) -> changelog
# 对每个段落按 RoleMap 取角色、从 FormatSpec 取规则，调用 core/executor.py 的
# 确定性函数改 XML；页边距/行网格走 section 级别。LLM 不碰 docx。

from docx import Document
from docx.shared import Mm

from core.executor import (
    set_doc_grid,
    set_run_fonts,
)
from core.style_set import (
    apply_named_style,
    clear_invalid_numbering_override,
    ensure_role_styles,
    resolve_target_body_style,
)
from core.track_changes import mark_paragraph_revision, snapshot_paragraph

_HF_ALIGN = {"left": 0, "center": 1, "right": 2, "justify": 3}

_ROLE_FALLBACKS = {
    "abstract_heading": ("heading_1", "body"),
    "abstract_body": ("body",),
    "keywords": ("body",),
    "chapter_heading": ("heading_1",),
    "bibliography_heading": ("heading_1",),
    "bibliography_entry": ("attachment", "body"),
    "equation": ("body",),
    "appendix_heading": ("heading_1",),
}


def _resolved_role(role, roles):
    if role in roles:
        return role
    return next((candidate for candidate in _ROLE_FALLBACKS.get(role, ()) if candidate in roles), None)


def _apply_header_footer(doc, page):
    """页眉页脚：text 写入（可选）+ 字体/字号/对齐；footer.page_number 插 PAGE 域。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    section = doc.sections[0]
    changed = []
    for which, container in (("header", section.header), ("footer", section.footer)):
        rule = page.get(which)
        if not isinstance(rule, dict) or not rule:
            continue
        if rule.get("text") is not None:
            # 覆盖第一段文字，字体随后统一刷
            p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = rule["text"]
            else:
                p.add_run(rule["text"])
            changed.append(f"{which}_text")
        if rule.get("page_number"):
            # 页码域：PAGE
            p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
            if p.runs or p.text:
                p = container.add_paragraph()
            run = p.add_run()
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = "PAGE"
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            run._element.append(fld_begin)
            run._element.append(instr)
            run._element.append(fld_end)
            changed.append(f"{which}_page_number")
        font_kwargs = {}
        if rule.get("font_eastasia"):
            font_kwargs["eastasia"] = rule["font_eastasia"]
        if rule.get("font_ascii"):
            font_kwargs["ascii_font"] = rule["font_ascii"]
        if rule.get("size_pt") is not None:
            font_kwargs["size_pt"] = rule["size_pt"]
        if rule.get("bold") is not None:
            font_kwargs["bold"] = rule["bold"]
        for p in container.paragraphs:
            if font_kwargs:
                for run in p.runs:
                    set_run_fonts(run, **font_kwargs)
            if rule.get("alignment") in _HF_ALIGN:
                p.alignment = _HF_ALIGN[rule["alignment"]]
        if font_kwargs:
            changed.append(f"{which}_font")
    return changed


def _apply_table_rule(doc, table_rule):
    """表格排版（v1）：首行表头加粗/居中 + 单元格字体字号 + 边框。
    只刷段落与字符格式，不动表格结构（合并单元格等保持原样）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not isinstance(table_rule, dict) or not table_rule:
        return []
    changed = False
    for table in doc.tables:
        changed = True
        if table_rule.get("borders"):
            tbl_pr = table._tbl.tblPr
            borders = tbl_pr.find(qn("w:tblBorders"))
            if borders is None:
                borders = OxmlElement("w:tblBorders")
                tbl_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = borders.find(qn(f"w:{edge}"))
                if el is None:
                    el = OxmlElement(f"w:{edge}")
                    borders.append(el)
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "4")
        for r_i, row in enumerate(table.rows):
            is_header = r_i == 0
            for cell in row.cells:
                for p in cell.paragraphs:
                    font_kwargs = {}
                    if table_rule.get("font_eastasia"):
                        font_kwargs["eastasia"] = table_rule["font_eastasia"]
                    if table_rule.get("font_ascii"):
                        font_kwargs["ascii_font"] = table_rule["font_ascii"]
                    if table_rule.get("size_pt") is not None:
                        font_kwargs["size_pt"] = table_rule["size_pt"]
                    if is_header and table_rule.get("header_bold"):
                        font_kwargs["bold"] = True
                    if font_kwargs:
                        for run in p.runs:
                            set_run_fonts(run, **font_kwargs)
                    align_key = "header_alignment" if is_header else "body_alignment"
                    if table_rule.get(align_key) in _HF_ALIGN:
                        p.alignment = _HF_ALIGN[table_rule[align_key]]
    return ["table_format"] if changed else []



def _apply_columns(doc, num_cols, before_element=None):
    """多栏排版：before_element（w:p 元素）之后的内容进入多栏（连续分节符切分）。
    before_element 为 None 时整篇多栏。"""
    import copy
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not isinstance(num_cols, int) or num_cols < 2:
        return False
    final_sect = doc.sections[0]._sectPr
    if before_element is not None:
        # 在切分点前的段落上插入单栏的段落级 sectPr（该段成为第一节的结尾）
        first_sect = copy.deepcopy(final_sect)
        cols = first_sect.find(qn("w:cols"))
        if cols is None:
            cols = OxmlElement("w:cols")
            doc_grid = first_sect.find(qn("w:docGrid"))
            if doc_grid is not None:
                first_sect.insert(list(first_sect).index(doc_grid), cols)
            else:
                first_sect.append(cols)
        cols.set(qn("w:num"), "1")
        ppr = before_element.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            before_element.insert(0, ppr)
        ppr.append(first_sect)
    cols = final_sect.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        doc_grid = final_sect.find(qn("w:docGrid"))
        if doc_grid is not None:
            final_sect.insert(list(final_sect).index(doc_grid), cols)
        else:
            final_sect.append(cols)
    cols.set(qn("w:num"), str(num_cols))
    return True


def _insert_toc(doc, before_paragraph, levels=(1, 2)):
    """在 before_paragraph 前插入目录标题 + TOC 域。Word 打开后更新域即生成目录。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _new_para():
        p = OxmlElement("w:p")
        before_paragraph._p.addprevious(p)
        return p

    # 目录标题段
    title_p = _new_para()
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "目录"
    r.append(t)
    title_p.append(r)

    # TOC 域
    lvl = "-".join(str(x) for x in sorted(levels))
    field_p = _new_para()

    def _fld(kind):
        el = OxmlElement("w:fldChar")
        el.set(qn("w:fldCharType"), kind)
        return el

    r1 = OxmlElement("w:r"); r1.append(_fld("begin"))
    r2 = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-%d" \\h \\z \\u' % max(levels)
    r2.append(instr)
    r3 = OxmlElement("w:r"); r3.append(_fld("separate"))
    r4 = OxmlElement("w:r")
    t4 = OxmlElement("w:t")
    t4.text = "（在 Word 中右键此处选择「更新域」生成目录）"
    r4.append(t4)
    r5 = OxmlElement("w:r"); r5.append(_fld("end"))
    for r in (r1, r2, r3, r4, r5):
        field_p.append(r)
    return True


def apply_format(docx_path, spec, rolemap, out_path, track=False,
                 template_path=None):
    """应用 FormatSpec × RoleMap，输出 docx，返回 changelog list[dict]。
    rolemap: {idx: role}（idx 对应 extract.py 的段落序号）。
    模板未明确指定的角色统一与正文保持一致（套用 body 规则与样式）。

    页面级：页边距/行网格 + 页眉页脚（page.header/footer，footer 支持页码域）；
    表格：spec.table 规则（首行表头加粗居中、单元格字体字号、边框）。
    track=True 时输出修订模式文档：段落/字符格式改动以 w:pPrChange /
    w:rPrChange 记录，Word 审阅视图可见。
    """
    doc = Document(docx_path)
    roles = spec.get("roles", {})
    cleanup_mode = (spec.get("cleanup") or {}).get("mode", "controlled")
    # 必须在创建/更新 FormatAgent 样式之前解析，避免把新样式误认成目标原样式。
    target_body_style = resolve_target_body_style(doc, rolemap)
    role_styles = ensure_role_styles(
        doc, spec, target_body_style=target_body_style)

    # ---- 页面级 ----
    page = spec.get("page") or {}
    margin = page.get("margin") or {}
    section = doc.sections[0]
    if margin.get("top_mm") is not None:
        section.top_margin = Mm(margin["top_mm"])
    if margin.get("bottom_mm") is not None:
        section.bottom_margin = Mm(margin["bottom_mm"])
    if margin.get("left_mm") is not None:
        section.left_margin = Mm(margin["left_mm"])
    if margin.get("right_mm") is not None:
        section.right_margin = Mm(margin["right_mm"])
    line_grid = page.get("line_grid") or {}
    if line_grid.get("line_pt") is not None:
        set_doc_grid(doc, line_pt=line_grid["line_pt"])
    # ---- 页眉页脚 + 表格 ----
    extra_changes = []
    extra_changes.extend(_apply_header_footer(doc, page))
    extra_changes.extend(_apply_table_rule(doc, spec.get("table")))
    # ---- 段落级 ----
    changelog = []
    rev_id = 1
    for idx, p in enumerate(doc.paragraphs):
        role = rolemap.get(idx, rolemap.get(str(idx)))
        if role is None:
            continue  # 未被标注的段落不动
        snapshot = snapshot_paragraph(p) if track else None
        resolved_role = _resolved_role(role, roles)
        if resolved_role is not None:
            rule = roles[resolved_role]
            style = role_styles[resolved_role]
            changed = apply_named_style(
                p, style, rule, role=role, cleanup_mode=cleanup_mode)
            fallback_to_target_body = resolved_role != role
        else:
            # 模板未规定的角色与正文保持一致：套用 body 的规则和命名样式，
            # 同时清掉会遮蔽样式的直接格式（如原文自带的加粗/异体字）。
            # 真实自动编号保留（body 不在清编号集合内），仅清“取消编号”残留。
            body_rule = roles.get("body", {})
            body_style = role_styles.get("body", target_body_style)
            changed = apply_named_style(
                p, body_style, body_rule, role="body", cleanup_mode=cleanup_mode)
            style = body_style
            fallback_to_target_body = True
        if track:
            rev_id = mark_paragraph_revision(p, snapshot, rev_id_start=rev_id)
        changelog.append({
            "idx": idx,
            "role": role,
            "style_name": style.name,
            "text": p.text.strip()[:30],
            "changed_fields": changed,
            "fallback_to_target_body": fallback_to_target_body,
        })

    # ---- 目录 + 多栏（结构级，在段落样式完成后执行）----
    toc = spec.get("toc") or {}
    columns = page.get("columns")
    structure_enabled = bool((spec.get("structure") or {}).get("enabled"))
    # 论文结构模式由 thesis_structure 生成独立目录节；不要再插一次旧版占位目录。
    want_toc = (
        isinstance(toc, dict) and toc.get("enabled") and not structure_enabled)
    want_cols = isinstance(columns, int) and columns >= 2
    if want_toc or want_cols:
        first_heading = None
        for idx, p in enumerate(doc.paragraphs):
            r = rolemap.get(idx, rolemap.get(str(idx)))
            if r in (
                "heading_1", "heading_2", "heading_3", "chapter_heading",
                "bibliography_heading", "appendix_heading",
            ):
                first_heading = p
                break
        if want_toc and first_heading is not None:
            # 目录要能收录我们的命名样式标题：强制补大纲级别
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn as _qn
            for role, lvl in (("heading_1", "0"), ("heading_2", "1"), ("heading_3", "2")):
                style = role_styles.get(role)
                if style is None:
                    continue
                ppr = style.element.get_or_add_pPr()
                if ppr.find(_qn("w:outlineLvl")) is None:
                    ol = OxmlElement("w:outlineLvl")
                    ol.set(_qn("w:val"), lvl)
                    ppr.append(ol)
            _insert_toc(doc, first_heading, levels=toc.get("levels") or (1, 2))
            extra_changes.append("toc_inserted")
        if want_cols:
            # 多栏从第一个标题开始（标题/副标题/摘要保持单栏，符合论文惯例）
            before_el = first_heading._p.getprevious() if first_heading is not None else None
            if _apply_columns(doc, columns, before_element=before_el):
                extra_changes.append(f"columns_{columns}")

    if extra_changes:
        changelog.append({
            "idx": -1,
            "role": "page",
            "style_name": "-",
            "text": "页眉页脚/表格/目录/多栏（页面级）",
            "changed_fields": extra_changes,
            "fallback_to_target_body": False,
        })

    # ---- 论文结构级：安全重建封面/前置页/目录/正文/参考文献 ----
    # 放在段落样式之后执行，确保 STYLEREF 与 TOC 能引用稳定的命名样式。
    if structure_enabled:
        from core.thesis_structure import assemble_thesis_structure
        structure_result = assemble_thesis_structure(
            doc, spec, rolemap, template_path=template_path)
        if structure_result:
            changelog.append({
                "idx": -2,
                "role": "structure",
                "style_name": "-",
                "text": "论文封面/前置页/目录/分节/动态页眉页码",
                "changed_fields": structure_result["changed_fields"],
                "fallback_to_target_body": False,
                "allowed_additions": structure_result["allowed_additions"],
                "stripped_prefixes": structure_result["stripped_prefixes"],
                "section_kinds": structure_result["section_kinds"],
            })

    doc.save(out_path)
    return changelog


def write_report(changelog, spec, report_path):
    """把 changelog 写成 markdown 修改对照报告。"""
    lines = ["# 排版修改对照报告", ""]
    page = spec.get("page") or {}
    if page:
        lines.append("## 页面设置")
        margin = page.get("margin") or {}
        if margin:
            lines.append(
                f"- 页边距（mm）：上 {margin.get('top_mm', '-')} / 下 {margin.get('bottom_mm', '-')}"
                f" / 左 {margin.get('left_mm', '-')} / 右 {margin.get('right_mm', '-')}")
        lg = page.get("line_grid") or {}
        if lg.get("line_pt"):
            lines.append(f"- 行网格：{lg['line_pt']} 磅/行")
        lines.append("")
    lines.append("## 段落修改明细")
    lines.append("")
    lines.append("| 段落 | 角色 | Word 样式 | 改动字段 | 内容摘要 |")
    lines.append("|---|---|---|---|---|")
    for c in changelog:
        fields = ", ".join(c["changed_fields"]) if c["changed_fields"] else "（无字段改动）"
        lines.append(
            f"| {c['idx']} | {c['role']} | {c.get('style_name', '-')} | {fields} | {c['text']} |")
    lines.append("")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return report_path
