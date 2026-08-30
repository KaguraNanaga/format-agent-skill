# 模板 docx → FormatSpec（格式源第二种，确定性读取，不依赖 VLM）。
# 思路：给定模板的 RoleMap，为每个角色找代表段落，用 effective_props 读生效
# 字体/字号/加粗，用 python-docx 读对齐/行距/缩进，页面级读 section 页边距和行网格。
# 未知/读不到的字段不编——留给 LLM 规范抽取或人肉 JSON 补。

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn

from core.effective_props import effective_props
from core.extract import manual_number_prefix, paragraph_numbering_metadata
from core.schema import validate_spec

_ALIGN_MAP = {0: "left", 1: "center", 2: "right", 3: "justify"}
_ALIGN_XML_MAP = {
    "left": "left",
    "start": "left",
    "center": "center",
    "right": "right",
    "end": "right",
    "both": "justify",
    "distribute": "justify",
    "thaiDistribute": "justify",
}


def _effective_ppr_elements(paragraph):
    """按“段落直接格式 → 段落样式链”顺序产出 pPr。"""
    direct = paragraph._p.pPr
    if direct is not None:
        yield direct
    style = paragraph.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        ppr = style.element.find(qn("w:pPr"))
        if ppr is not None:
            yield ppr
        style = style.base_style


def _effective_ppr_child(paragraph, tag):
    for ppr in _effective_ppr_elements(paragraph):
        child = ppr.find(qn(tag))
        if child is not None:
            return child
    return None


def _effective_ppr_attr(paragraph, child_tag, *attrs):
    """逐层、逐属性读取有效值，允许直接格式只覆盖同一元素的部分属性。"""
    for ppr in _effective_ppr_elements(paragraph):
        child = ppr.find(qn(child_tag))
        if child is None:
            continue
        for attr in attrs:
            value = child.get(qn(attr))
            if value is not None:
                return value
    return None


def _numbering_element_for_paragraph(p):
    """返回段落直接或样式链继承的 numPr。"""
    ppr = p._p.pPr
    if ppr is not None:
        num_pr = ppr.find(qn("w:numPr"))
        if num_pr is not None:
            return num_pr
    style = p.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        ppr = style.element.find(qn("w:pPr"))
        if ppr is not None:
            num_pr = ppr.find(qn("w:numPr"))
            if num_pr is not None:
                return num_pr
        style = style.base_style
    return None


def _find_by_attr(parent, tag, attr, value):
    for element in parent.findall(qn(tag)):
        if element.get(qn(attr)) == str(value):
            return element
    return None


def _val(parent, tag, default=None):
    element = parent.find(qn(tag)) if parent is not None else None
    return element.get(qn("w:val")) if element is not None else default


def _int_or_none(value):
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _numbering_rule(doc, paragraph, role):
    """解析段落的有效 OOXML 编号级别，转成可跨文档复制的规则。"""
    # 正文的自动列表是局部段落结构，不是“所有正文”的样式规则。
    # 若把偶然命中的第一个 body 列表项抽成 body.numbering，会使全文正文
    # 都被编号。当前只允许语义标题角色产生全局样式编号。
    if not role.startswith("heading_"):
        return None
    num_pr = _numbering_element_for_paragraph(paragraph)
    if num_pr is None:
        return None
    num_id = _int_or_none(_val(num_pr, "w:numId"))
    level = _int_or_none(_val(num_pr, "w:ilvl", "0"))
    if num_id is None or num_id <= 0 or level is None:
        return None

    numbering = doc.part.numbering_part.element
    num = _find_by_attr(numbering, "w:num", "w:numId", num_id)
    if num is None:
        return None
    abstract_id = _int_or_none(_val(num, "w:abstractNumId"))
    abstract = _find_by_attr(
        numbering, "w:abstractNum", "w:abstractNumId", abstract_id)
    if abstract is None:
        return None

    lvl = None
    override = _find_by_attr(num, "w:lvlOverride", "w:ilvl", level)
    if override is not None:
        lvl = override.find(qn("w:lvl"))
    if lvl is None:
        lvl = _find_by_attr(abstract, "w:lvl", "w:ilvl", level)
    if lvl is None:
        return None

    num_format = _val(lvl, "w:numFmt")
    level_text = _val(lvl, "w:lvlText")
    if not num_format or level_text is None:
        return None
    override_start = _int_or_none(_val(override, "w:startOverride")) if override is not None else None
    rule = {
        "group": "headings" if role.startswith("heading_") else role,
        "level": level,
        "num_format": num_format,
        "level_text": level_text,
        "start": override_start or _int_or_none(_val(lvl, "w:start", "1")) or 1,
        "suffix": _val(lvl, "w:suff", "tab"),
        "alignment": _val(lvl, "w:lvlJc", "left"),
    }
    if lvl.find(qn("w:isLgl")) is not None:
        rule["is_legal"] = True
    level_restart = _int_or_none(_val(lvl, "w:lvlRestart"))
    if level_restart is not None:
        rule["level_restart"] = level_restart

    ppr = lvl.find(qn("w:pPr"))
    ind = ppr.find(qn("w:ind")) if ppr is not None else None
    if ind is not None:
        left = ind.get(qn("w:left")) or ind.get(qn("w:start"))
        for value, key in (
            (left, "left_twips"),
            (ind.get(qn("w:hanging")), "hanging_twips"),
            (ind.get(qn("w:firstLine")), "first_line_twips"),
        ):
            parsed = _int_or_none(value)
            if parsed is not None:
                rule[key] = parsed
    tabs = ppr.find(qn("w:tabs")) if ppr is not None else None
    if tabs is not None:
        for tab in tabs.findall(qn("w:tab")):
            if tab.get(qn("w:val")) in {"num", "left"}:
                value = _int_or_none(tab.get(qn("w:pos")))
                if value is not None:
                    rule["tab_pos_twips"] = value
                    break
    rpr = lvl.find(qn("w:rPr"))
    if rpr is not None:
        fonts = rpr.find(qn("w:rFonts"))
        if fonts is not None:
            eastasia = fonts.get(qn("w:eastAsia"))
            ascii_font = fonts.get(qn("w:ascii")) or fonts.get(qn("w:hAnsi"))
            if eastasia:
                rule["font_eastasia"] = eastasia
            if ascii_font:
                rule["font_ascii"] = ascii_font
        size = _int_or_none(_val(rpr, "w:sz"))
        if size is not None:
            rule["size_pt"] = size / 2
        bold = rpr.find(qn("w:b"))
        if bold is not None:
            rule["bold"] = bold.get(qn("w:val"), "1") not in {"0", "false", "off"}
    return rule


def _is_numbered_body_candidate(paragraph):
    """用于选正文代表段：优先避开真自动列表和手工 1./1.2 列表。"""
    metadata = paragraph_numbering_metadata(paragraph)
    return bool(
        metadata.get("numbering_status") == "automatic"
        or manual_number_prefix(paragraph.text) is not None
    )


def _representative_paragraphs(paragraphs, rolemap):
    """按角色选代表段；body 优先选普通正文，避免列表特例污染全局规则。"""
    candidates = {}
    for idx, role in sorted(rolemap.items()):
        if 0 <= idx < len(paragraphs):
            candidates.setdefault(role, []).append(paragraphs[idx])
    representatives = {}
    for role, role_paragraphs in candidates.items():
        if role == "body":
            preferred = [
                paragraph for paragraph in role_paragraphs
                if paragraph.text.strip() and not _is_numbered_body_candidate(paragraph)
            ]
            representatives[role] = (preferred or role_paragraphs)[0]
        else:
            representatives[role] = role_paragraphs[0]
    return representatives


def _para_alignment(p):
    jc = _effective_ppr_child(p, "w:jc")
    if jc is not None:
        value = jc.get(qn("w:val"))
        if value in _ALIGN_XML_MAP:
            return _ALIGN_XML_MAP[value]
    a = p.alignment
    return _ALIGN_MAP.get(int(a)) if a is not None else None


def _para_line_spacing(p):
    spacing = _effective_ppr_child(p, "w:spacing")
    if spacing is not None and spacing.get(qn("w:line")) is not None:
        line = _int_or_none(spacing.get(qn("w:line")))
        line_rule = spacing.get(qn("w:lineRule"), "auto")
        if line is not None and line_rule == "exact":
            return {"type": "exact", "pt": round(line / 20, 1)}
        if line is not None and line_rule == "auto":
            return {"type": "multiple", "pt": round(line / 240, 2)}
    pf = p.paragraph_format
    rule = pf.line_spacing_rule
    if rule == WD_LINE_SPACING.EXACTLY and pf.line_spacing is not None:
        return {"type": "exact", "pt": round(pf.line_spacing.pt, 1)}
    if rule == WD_LINE_SPACING.MULTIPLE and pf.line_spacing is not None:
        return {"type": "multiple", "pt": round(float(pf.line_spacing), 2)}
    return None


def _para_indent_chars(p, size_pt):
    """首行缩进字符数：优先读 XML firstLineChars，否则用磅值/字号反推。"""
    # hanging 与 firstLine 互斥；较近一层明确设置悬挂时，不能继续继承
    # 样式中的首行缩进。
    for ppr in _effective_ppr_elements(p):
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            continue
        if ind.get(qn("w:hanging")) is not None or ind.get(qn("w:hangingChars")) is not None:
            return None
        flc = ind.get(qn("w:firstLineChars"))
        if flc is not None:
            return round(int(flc) / 100, 1)
        first_line = ind.get(qn("w:firstLine"))
        if first_line is not None and size_pt:
            return round((int(first_line) / 20) / size_pt, 1)
    fl = p.paragraph_format.first_line_indent
    if fl is not None and size_pt:
        return round(fl.pt / size_pt, 1)
    return None


def _para_spacing_pt(p, attr):
    value = _effective_ppr_attr(p, "w:spacing", attr)
    parsed = _int_or_none(value)
    return round(parsed / 20, 1) if parsed is not None else None


def _page_section(doc):
    page = {}
    s = doc.sections[0]
    margin = {
        "top_mm": round(s.top_margin.mm, 1),
        "bottom_mm": round(s.bottom_margin.mm, 1),
        "left_mm": round(s.left_margin.mm, 1),
        "right_mm": round(s.right_margin.mm, 1),
    }
    if all(v is not None for v in margin.values()):
        page["margin"] = margin
    doc_grid = s._sectPr.find(qn("w:docGrid"))
    if doc_grid is not None and doc_grid.get(qn("w:linePitch")):
        page["line_grid"] = {"line_pt": round(int(doc_grid.get(qn("w:linePitch"))) / 20, 1)}
    # 多栏（论文双栏等）
    cols = s._sectPr.find(qn("w:cols"))
    if cols is not None and cols.get(qn("w:num")):
        try:
            num = int(cols.get(qn("w:num")))
            if num >= 2:
                page["columns"] = num
        except ValueError:
            pass
    return page


def _has_toc(doc):
    """模板里是否有目录域（TOC field）。"""
    body = doc.element.body
    for el in body.iter(qn("w:instrText")):
        if "TOC" in (el.text or ""):
            return True
    return False


def _header_footer_rules(doc):
    """读模板的页眉页脚：第一个非空段的文字 + 生效格式；页脚有 PAGE 域则记 page_number。"""
    rules = {}
    s = doc.sections[0]
    for which, hf in (("header", s.header), ("footer", s.footer)):
        try:
            text_ps = [p for p in hf.paragraphs if p.text.strip()]
        except Exception:  # 无页眉页脚部件
            continue
        has_page_field = any(
            "PAGE" in (el.text or "")
            for el in hf._element.iter(qn("w:instrText")))
        if not text_ps and not has_page_field:
            continue
        rule = {}
        if text_ps:
            p = text_ps[0]
            rule["text"] = p.text.strip()
            props = effective_props(p)
            if props.get("eastasia"):
                rule["font_eastasia"] = props["eastasia"]
            if props.get("ascii"):
                rule["font_ascii"] = props["ascii"]
            if props.get("size_pt"):
                rule["size_pt"] = props["size_pt"]
            if props.get("bold") is not None:
                rule["bold"] = bool(props["bold"])
            a = _para_alignment(p)
            if a:
                rule["alignment"] = a
        if which == "footer" and has_page_field:
            rule["page_number"] = True
        rules[which] = rule
    return rules


def _table_rule(doc):
    """读模板第一张表：表头行格式（加粗/居中）+ 单元格字体字号 + 是否有边框。"""
    if not doc.tables:
        return None
    t = doc.tables[0]
    rule = {}
    rows = t.rows
    if not rows:
        return None
    header_props = None
    for p in rows[0].cells[0].paragraphs:
        header_props = effective_props(p)
        a = _para_alignment(p)
        if a:
            rule["header_alignment"] = a
        break
    if header_props and header_props.get("bold") is not None:
        rule["header_bold"] = bool(header_props["bold"])
    # 正文行字体字号
    body_row = rows[1] if len(rows) > 1 else rows[0]
    for p in body_row.cells[0].paragraphs:
        props = effective_props(p)
        if props.get("eastasia"):
            rule["font_eastasia"] = props["eastasia"]
        if props.get("ascii"):
            rule["font_ascii"] = props["ascii"]
        if props.get("size_pt"):
            rule["size_pt"] = props["size_pt"]
        a = _para_alignment(p)
        if a:
            rule["body_alignment"] = a
        break
    if header_props:
        rule.setdefault("font_eastasia", header_props.get("eastasia") or "宋体")
        rule.setdefault("size_pt", header_props.get("size_pt") or 10.5)

    # 边框：直接 tblBorders 或表格样式（Table Grid 等）继承都算有边框
    borders = t._tbl.tblPr.find(qn("w:tblBorders"))
    if borders is not None:
        any_border = any(
            (el.get(qn("w:val")) or "single") not in ("none", "nil")
            for el in borders)
        rule["borders"] = any_border
    else:
        style_name = ""
        try:
            style_name = (t.style.name or "").lower() if t.style else ""
        except Exception:
            style_name = ""
        if "grid" in style_name or "网格" in style_name:
            rule["borders"] = True
    return rule or None


def extract_rules_from_template(template_path, rolemap):
    """模板 docx + RoleMap → FormatSpec（经 schema 校验）。
    rolemap: {idx: role}。每个角色取第一个代表段落读格式。
    要求 rolemap 里至少有 body 角色，否则抛 ValueError。
    """
    doc = Document(template_path)
    paras = doc.paragraphs
    roles = {}
    representatives = _representative_paragraphs(paras, rolemap)
    for role, p in representatives.items():
        props = effective_props(p)
        eastasia = props.get("eastasia") or "宋体"
        size_pt = props.get("size_pt") or 10.5
        rule = {"font_eastasia": eastasia, "size_pt": size_pt,
                "bold": bool(props.get("bold"))}
        if props.get("ascii"):
            rule["font_ascii"] = props["ascii"]
        a = _para_alignment(p)
        if a:
            rule["alignment"] = a
        ls = _para_line_spacing(p)
        if ls:
            rule["line_spacing"] = ls
        flc = _para_indent_chars(p, size_pt)
        if flc:
            rule["first_line_indent_chars"] = flc
        space_before = _para_spacing_pt(p, "w:before")
        space_after = _para_spacing_pt(p, "w:after")
        if space_before is not None:
            rule["space_before_pt"] = space_before
        if space_after is not None:
            rule["space_after_pt"] = space_after
        numbering = _numbering_rule(doc, p, role)
        if numbering is not None:
            rule["numbering"] = numbering
        roles[role] = rule

    if "body" not in roles:
        raise ValueError("模板中没有标注 body 角色的段落，无法确定正文格式")
    # schema 要求每个角色至少有 alignment；读不到时给合理默认
    for role, rule in roles.items():
        rule.setdefault("alignment", "justify" if role == "body" else "left")

    spec = {"page": _page_section(doc), "roles": roles}
    # 页眉页脚 + 表格规则（模板有就读，没有就不编）
    spec["page"].update(_header_footer_rules(doc))
    table_rule = _table_rule(doc)
    if table_rule:
        spec["table"] = table_rule
    if _has_toc(doc):
        spec["toc"] = {"enabled": True, "levels": [1, 2]}
    # 行网格一致性：模板里的 docGrid 常是 Word 默认值（如 15.6pt），与正文实际
    # 固定行距不一致时，网格会干扰排版。正文有明确固定行距时，以正文行距为准。
    body_ls = (roles.get("body") or {}).get("line_spacing") or {}
    if body_ls.get("type") == "exact" and body_ls.get("pt"):
        grid = spec["page"].setdefault("line_grid", {})
        if grid.get("line_pt") != body_ls["pt"]:
            grid["line_pt"] = body_ls["pt"]
    validate_spec(spec)
    return spec


if __name__ == "__main__":
    import json
    import sys
    with open(sys.argv[2], encoding="utf-8") as f:
        rolemap = {int(k): v for k, v in json.load(f).items()}
    spec = extract_rules_from_template(sys.argv[1], rolemap)
    print(json.dumps(spec, ensure_ascii=False, indent=2))
