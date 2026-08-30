# docx → 段落清单（PLAN.md 第 7 节）。
# 除基础格式外，同时输出结构化编号、大纲和缩进元数据。
# size_pt/bold 用 effective_props 读"生效属性"而非样式名。
# in_table=True 的段落 v1 不参与重排，但仍列出（供角色标注参考上下文）。

import re

from docx import Document
from docx.oxml.ns import qn

from core.effective_props import get_paragraph_effective_font

_ALIGN_MAP = {0: "left", 1: "center", 2: "right", 3: "justify"}

_SENTENCE_ENDINGS = ("。", "；", ";", "！", "!", "？", "?", "，", ",", "：", ":")

# 先匹配 1.2 / 1.2.1，再匹配 1. / 2、，避免把 1.2 误拆成 "1." 。
# 层级数字后必须有空白或标点，因此“1.2万元”不会被当成列表前缀。
_MANUAL_HIERARCHICAL = re.compile(
    r"^\s*(?P<label>\d+(?:\.\d+)+)(?:[.、．)](?=\s|[\u4e00-\u9fff])|(?=\s))\s*")
_MANUAL_SIMPLE = re.compile(
    r"^\s*(?P<label>\d{1,3})[.、．)](?=\s|[\u4e00-\u9fff])\s*")


def _int_or_none(value):
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _val(parent, tag, default=None):
    element = parent.find(qn(tag)) if parent is not None else None
    return element.get(qn("w:val")) if element is not None else default


def _find_by_attr(parent, tag, attr, value):
    if parent is None:
        return None
    for element in parent.findall(qn(tag)):
        if element.get(qn(attr)) == str(value):
            return element
    return None


def manual_number_prefix(text):
    """返回手工数字前缀元数据；无明确前缀时返回 None。

    支持 ``1.``、``2、``、``1.2`` 和 ``1.2.1``。这只说明文本里有
    手工标记，不等于该段是标题。
    """
    value = (text or "").strip()
    match = _MANUAL_HIERARCHICAL.match(value) or _MANUAL_SIMPLE.match(value)
    if match is None:
        return None
    label = match.group("label")
    try:
        parts = [int(part) for part in label.split(".")]
    except ValueError:
        return None
    return {"label": label, "parts": parts, "prefix": match.group(0)}


def _style_chain(paragraph):
    style = paragraph.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        yield style
        style = style.base_style


def _effective_element(paragraph, tag):
    """返回段落直刷或样式链中首个指定元素及来源。"""
    ppr = paragraph._p.pPr
    direct = ppr.find(qn(tag)) if ppr is not None else None
    if direct is not None:
        return direct, "direct"
    for style in _style_chain(paragraph):
        style_ppr = style.element.find(qn("w:pPr"))
        element = style_ppr.find(qn(tag)) if style_ppr is not None else None
        if element is not None:
            return element, "style"
    return None, None


def paragraph_numbering_metadata(paragraph):
    """解析段落的有效 OOXML 编号态。

    ``automatic`` 必须同时满足 numId > 0 且能在 numbering.xml 中
    解析出定义。numId=0 或 ilvl<0 是取消编号，不得冒充自动列表。
    """
    num_pr, source = _effective_element(paragraph, "w:numPr")
    base = {
        "numbering_status": "none",
        "numbering_source": source,
        "num_id": None,
        "num_level": None,
        "num_format": None,
        "level_text": None,
        "numbering_left_twips": None,
        "numbering_hanging_twips": None,
        "numbering_first_line_twips": None,
    }
    if num_pr is None:
        return base

    num_id = _int_or_none(_val(num_pr, "w:numId"))
    level = _int_or_none(_val(num_pr, "w:ilvl", "0"))
    base.update({"num_id": num_id, "num_level": level})
    if num_id is None or num_id <= 0 or level is None or level < 0:
        base["numbering_status"] = "cancelled"
        return base

    try:
        numbering = paragraph.part.numbering_part.element
    except (AttributeError, KeyError):
        base["numbering_status"] = "unresolved"
        return base
    num = _find_by_attr(numbering, "w:num", "w:numId", num_id)
    abstract_id = _int_or_none(_val(num, "w:abstractNumId")) if num is not None else None
    abstract = _find_by_attr(numbering, "w:abstractNum", "w:abstractNumId", abstract_id)
    if num is None or abstract is None:
        base["numbering_status"] = "unresolved"
        return base

    override = _find_by_attr(num, "w:lvlOverride", "w:ilvl", level)
    lvl = override.find(qn("w:lvl")) if override is not None else None
    if lvl is None:
        lvl = _find_by_attr(abstract, "w:lvl", "w:ilvl", level)
    if lvl is None:
        base["numbering_status"] = "unresolved"
        return base

    num_format = _val(lvl, "w:numFmt")
    level_text = _val(lvl, "w:lvlText")
    if not num_format or level_text is None:
        base["numbering_status"] = "unresolved"
        return base
    base.update({
        "numbering_status": "automatic",
        "num_format": num_format,
        "level_text": level_text,
    })
    level_ppr = lvl.find(qn("w:pPr"))
    ind = level_ppr.find(qn("w:ind")) if level_ppr is not None else None
    if ind is not None:
        left = ind.get(qn("w:left")) or ind.get(qn("w:start"))
        base.update({
            "numbering_left_twips": _int_or_none(left),
            "numbering_hanging_twips": _int_or_none(ind.get(qn("w:hanging"))),
            "numbering_first_line_twips": _int_or_none(ind.get(qn("w:firstLine"))),
        })
    return base


def _outline_level(paragraph):
    outline, _ = _effective_element(paragraph, "w:outlineLvl")
    value = _int_or_none(outline.get(qn("w:val"))) if outline is not None else None
    return value if value is not None and 0 <= value <= 8 else None


def _direct_indent(paragraph):
    ppr = paragraph._p.pPr
    ind = ppr.find(qn("w:ind")) if ppr is not None else None
    if ind is None:
        return {
            "indent_left_twips": None,
            "indent_hanging_twips": None,
            "indent_first_line_twips": None,
            "indent_first_line_chars": None,
        }
    return {
        "indent_left_twips": _int_or_none(ind.get(qn("w:left"))),
        "indent_hanging_twips": _int_or_none(ind.get(qn("w:hanging"))),
        "indent_first_line_twips": _int_or_none(ind.get(qn("w:firstLine"))),
        "indent_first_line_chars": _int_or_none(ind.get(qn("w:firstLineChars"))),
    }


def _alignment_name(paragraph):
    a = paragraph.alignment
    if a is None:
        return None
    return _ALIGN_MAP.get(int(a), str(a))


def _spacing_pt(paragraph):
    """读段落直刷的段前/段后距（磅）。样式层定义的间距读不到，返回 None。"""
    pf = paragraph.paragraph_format
    sb = round(pf.space_before.pt, 1) if pf.space_before is not None else None
    sa = round(pf.space_after.pt, 1) if pf.space_after is not None else None
    return sb, sa


def _para_record(idx, p, in_table):
    eastasia, size_pt, bold = get_paragraph_effective_font(p)
    sb, sa = _spacing_pt(p)
    full_text = p.text.strip()
    manual = manual_number_prefix(full_text)
    numbering = paragraph_numbering_metadata(p)
    if numbering["numbering_status"] == "automatic":
        list_kind = "automatic"
    elif manual is not None:
        list_kind = "manual"
    else:
        list_kind = "none"
    record = {
        "idx": idx,
        "text": full_text[:80],
        "char_count": len(full_text),
        "ends_with_sentence_punct": full_text.endswith(_SENTENCE_ENDINGS),
        "size_pt": size_pt,
        "bold": bold,
        "alignment": _alignment_name(p),
        "style_name": p.style.name if p.style is not None else None,
        "outline_level": _outline_level(p),
        "list_kind": list_kind,
        "manual_number": manual["label"] if manual is not None else None,
        "manual_number_parts": manual["parts"] if manual is not None else None,
        "list_sequence": False,
        "space_before_pt": sb,
        "space_after_pt": sa,
        "in_table": in_table,
    }
    record.update(numbering)
    record.update(_direct_indent(p))
    return record


def _manual_numbers_are_consecutive(previous, current):
    a = previous.get("manual_number_parts")
    b = current.get("manual_number_parts")
    return bool(
        isinstance(a, list) and isinstance(b, list) and len(a) == len(b)
        and a[:-1] == b[:-1] and b[-1] == a[-1] + 1
    )


def _automatic_numbers_share_sequence(previous, current):
    return bool(
        previous.get("numbering_status") == "automatic"
        and current.get("numbering_status") == "automatic"
        and previous.get("num_id") == current.get("num_id")
        and previous.get("num_level") == current.get("num_level")
    )


def _annotate_list_sequences(records):
    """标记相邻的 1/2/3 或 1.1/1.2 手工序列，以及同 numId/ilvl
    的真自动列表。空段、表格边界和非列表段会中断序列。
    """
    for previous, current in zip(records, records[1:]):
        if previous.get("in_table") != current.get("in_table"):
            continue
        if not previous.get("text") or not current.get("text"):
            continue
        if (
            _manual_numbers_are_consecutive(previous, current)
            or _automatic_numbers_share_sequence(previous, current)
        ):
            previous["list_sequence"] = True
            current["list_sequence"] = True


def extract_paragraphs(docx_path):
    """返回段落清单 list[dict]，idx 为全文段落序号（含表格内段落）。"""
    doc = Document(docx_path)
    out = []
    for idx, p in enumerate(doc.paragraphs):
        out.append(_para_record(idx, p, False))
    # 表格段落追加在末尾（doc.paragraphs 不含表格内段落，单独列出）
    idx = len(out)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    out.append(_para_record(idx, p, True))
                    idx += 1
    _annotate_list_sequences(out)
    return out


if __name__ == "__main__":
    import json
    import sys
    paras = extract_paragraphs(sys.argv[1])
    print(json.dumps(paras, ensure_ascii=False, indent=2))
