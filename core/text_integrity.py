# 文本一致性校验：排版只许改格式，不许动文字。
# check_text_integrity(原稿, 终稿) -> 逐段比对结果。
# 设计要点：
# - 规范化只处理空白（Word 渲染产生的多余空格），不做任何语义改写
# - 两种"合法差异"不算失败：
#   1) 我们插入的目录占位文字（allowed_additions，由调用方告知）
#   2) 手工编号被自动编号替换而剥掉的前缀（expected_stripped_prefixes）
# - 其余任何文字增删改都会报告出来

import difflib
import re
from collections import Counter

from docx import Document


def _norm(text):
    """空白规范化：连续空白合成一个，去首尾。"""
    return re.sub(r"\s+", " ", text or "").strip()


def paragraph_texts(docx_path):
    """提取文档全部段落（含表格内）的规范化文字，跳过空段。"""
    doc = Document(docx_path)
    texts = [_norm(p.text) for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(_norm(p.text) for p in cell.paragraphs)
    return [t for t in texts if t]


def check_text_integrity(source_path, out_path,
                         allowed_additions=(), expected_stripped_prefixes=()):
    """比对原稿与终稿的段落文字。
    allowed_additions: 允许出现的新增段落文字（如目录标题与占位提示）。
    expected_stripped_prefixes: 允许被剥掉的手工编号前缀（如 ["1.", "2."]）。
    返回 {"ok": bool, "added": [...], "removed": [...], "changed": [...]}。
    """
    src = paragraph_texts(source_path)
    out = paragraph_texts(out_path)

    # 剥掉合法新增段
    # Counter 而不是 set：封面多个元数据字段可能合法地使用同一占位值。
    allowed = Counter(_norm(a) for a in allowed_additions)
    out_filtered = []
    for t in out:
        if allowed[t] > 0:
            allowed[t] -= 1  # 每个声明的合法新增只抵消一次
            continue
        out_filtered.append(t)

    # 剥掉"手工编号前缀被自动编号替换"的段落差异：
    # 终稿段落 = 原稿段落去掉前缀，视为等价，两边同时归一
    prefixes = sorted(expected_stripped_prefixes, key=len, reverse=True)

    def strip_known_prefix(t):
        for pre in prefixes:
            if t.startswith(pre):
                return _norm(t[len(pre):])
        return t

    if prefixes:
        src = [strip_known_prefix(t) for t in src]
        out_filtered = [strip_known_prefix(t) for t in out_filtered]

    added, removed, changed = [], [], []
    matcher = difflib.SequenceMatcher(a=src, b=out_filtered, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        if tag in ("replace", "delete"):
            removed.extend(src[i1:i2])
        if tag in ("replace", "insert"):
            added.extend(out_filtered[j1:j2])
        if tag == "replace" and (i2 - i1) == (j2 - j1):
            changed = [{"from": a, "to": b} for a, b in zip(src[i1:i2], out_filtered[j1:j2])]

    ok = not added and not removed
    return {"ok": ok, "added": added, "removed": removed, "changed": changed}


if __name__ == "__main__":
    import json
    import sys
    print(json.dumps(check_text_integrity(sys.argv[1], sys.argv[2]),
                     ensure_ascii=False, indent=2))
